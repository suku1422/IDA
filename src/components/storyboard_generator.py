from datetime import datetime
import streamlit as st
import pandas as pd

from src.openai_client import get_openai_response

def generate_storyboard():
    # if not context_summary or not content_outline:
    #     raise ValueError("Context summary and content outline are required to generate a storyboard.")

    # prompt = (
    #     f"Create a storyboard that follows instructional design theories for the e-learning course based on the following instructional design context, content outline, and source content.\n\n"
    #     f"### Instructional Design Context:\n{context_summary}\n\n"
    #     f"### Content Outline:\n{content_outline}\n\n"
    #     f"### Source Content:\n{source_content}\n\n"
    #     f"Provide the storyboard as a table with three columns: Onscreen Text | Voice Over Script | Visualization Guidelines.\n"
    #     f"Start immediately with the table header. Separate columns using a '|' (pipe symbol).\n"
    #     f"Make sure that the Onscreen text column contains the entire text we want to include in the slide, not just slide titles. In the Voice over script column, include the entire narrative voice over script, not just an introduction.\n"
    #     f"Let knowledge checks not be too many. Also, when knowledge checks are used, include all the details - the question, the answer options, the correct answer, and also correct and wrong answer feedback.\n"
    #     f"Do not add any explanation before or after the table. Each row must be properly formatted without bullets or other formatting."
    # )

    # # Call to OpenAI API would go here to generate the storyboard based on the prompt
    # storyboard = get_openai_response(prompt)
    
    # return storyboard

    st.header("Step 4: Generate Storyboard")
    context_summary = st.session_state.get("context_summary_persisted", "")
    content_outline = st.session_state.get("content_outline", "")
    uploaded_content = st.session_state.get("uploaded_content", "")
    generated_content = st.session_state.get("generated_additional_content", "")

    if not content_outline:
        st.error("No content outline found. Please complete Step 3 before proceeding.")
        return

    combined_content = uploaded_content.strip()
    if generated_content:
        combined_content += f"\n\n{generated_content.strip()}"

    if st.button("🔁 Regenerate Storyboard"):
        st.session_state.storyboard = None  # Reset
                        
    if not st.session_state.get("storyboard"):
        prompt = (
            f"Create a storyboard that follows instructional design theories, for the e-learning course based on the following instructional design context, content outline, and source content.\n\n"
            f"### Instructional Design Context:\n{context_summary}\n\n"
            f"### Content Outline:\n{content_outline}\n\n"
            f"### Source Content:\n{combined_content}\n\n"
            f"Provide the storyboard as a table with three columns: Onscreen Text | Voice Over Script | Visualization Guidelines.\n"
            f"Start immediately with the table header. Separate columns using a '|' (pipe symbol).\n"
            f"Make sure that the Onscreen text column contains the entire text we want to include in the slide, not just slide titles. In the Voice over script column, include the entire narrative voice over script, not just an introduction.\n"
            f"Let knowledge checks not be too many. Also when knowledge checks are used include all the details - the question, the answer options, the correct answer and also correct and wrong answer feedback. \n"
            f"Do not add any explanation before or after the table. Each row must be properly formatted without bullets or other formatting."
        )

        with st.spinner("Generating storyboard. This may take a while..."):
            storyboard = get_openai_response(prompt, max_completion_tokens=16384)
            if storyboard:
                st.session_state.storyboard = storyboard

    storyboard_text = st.session_state.get("storyboard", "")

    if storyboard_text:
        st.subheader("Generated Storyboard")
        import io
        import pandas as pd
        try:
            df = pd.read_csv(io.StringIO(st.session_state.storyboard), sep="|")
            df.columns = [col.strip() for col in df.columns]
            df = df.loc[:, ~df.columns.str.contains("^Unnamed")]
            df = df.reset_index(drop=True)

            if df.shape[1] > 3:
                df = df.iloc[:, -3:]

            df.columns = ["Onscreen Text", "Voice Over Script", "Visualization Guidelines"]

            st.markdown(
                """
                <style>
                .stDataFrame td {
                    white-space: pre-wrap !important;
                    word-break: break-word !important;
                }
                </style>
                """,
                unsafe_allow_html=True
            )

            st.dataframe(df, use_container_width=True)

        except Exception:
            st.error("Could not render storyboard as a table here. Showing raw text instead. Don't worry, export to word and it will be fine")
            st.code(st.session_state.storyboard)

        # Show "Continue to Step 5" only if storyboard exists
        with st.form("approve_storyboard_form"):
            submitted = st.form_submit_button("✅ Continue to Step 5")
            if submitted:
                st.session_state.step = 5
                st.rerun()
            
        # Export to Word
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import io

        lines = storyboard_text.strip().split("\n")

        if len(lines) >= 2:
            buffer = io.BytesIO()
            doc = Document()
            doc.add_heading('Storyboard', level=1)

            headers = [h.strip() for h in lines[0].split("|") if h.strip()]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            table.autofit = True
            hdr_cells = table.rows[0].cells
            for idx, col in enumerate(headers):
                cell = hdr_cells[idx]
                cell.text = col
                for paragraph in cell.paragraphs:
                    run = paragraph.runs[0]
                    run.bold = True
                    run.font.size = Pt(11)

            for line in lines[2:]:
                columns = [c.strip() for c in line.split("|") if c.strip()]
                if len(columns) == len(headers):
                    row_cells = table.add_row().cells
                    for idx, val in enumerate(columns):
                        cell = row_cells[idx]
                        paragraph = cell.paragraphs[0]
                        run = paragraph.add_run(str(val))
                        run.font.size = Pt(10)

                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcW = OxmlElement('w:tcW')
                        tcW.set(qn('w:type'), 'auto')
                        tcW.set(qn('w:w'), '0')
                        tcPr.append(tcW)

            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="\U0001F4C4 Download Storyboard as Word file",
                data=buffer,
                file_name="Storyboard.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.error("Storyboard format not valid. Cannot export to Word.")
    else:
        st.error("Storyboard could not be generated. Please retry.")

    # with st.form("approve_storyboard_form"):
    #     submitted = st.form_submit_button("Approve Storyboard and Continue", type="primary")
    #     if submitted:
    #         st.session_state.step = 5
    #         st.rerun()


def save_storyboard_to_file(storyboard, file_path):
    with open(file_path, 'w') as file:
        file.write(storyboard)
    print(f"Storyboard saved to {file_path}")

def load_storyboard_from_file(file_path):
    with open(file_path, 'r') as file:
        return file.read()