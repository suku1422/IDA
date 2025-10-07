import re
import streamlit as st
from src.openai_client import get_openai_response


def create_final_assessment(context_summary = st.session_state.get("context_summary", ""), content_outline = st.session_state.get("content_outline", ""), num_questions = 5):
    # prompt = (
    #     f"Based on the following instructional design context and content outline, generate a final assessment for this e-learning course.\n\n"
    #     f"### Instructional Design Context:\n{context_summary}\n\n"
    #     f"### Content Outline:\n{content_outline}\n\n"
    #     f"Create {num_questions} multiple-choice questions.\n"
    #     f"Each MCQ must have appropriate number of answer options, and should clearly indicate the correct option.\n"
    #     f"Ensure questions align with the course objectives and learning content.\n"
    #     f"Do not add any explanation text or headings before or after the questions."
    # )
    
    # assessment = get_openai_response(prompt, max_completion_tokens=4500)
    # return assessment
    if st.session_state.get("assessment_downloaded"):
        st.success("🎉 Instructional design process completed successfully!")
        return
    
    st.header("Step 5: Create Final Assessment")

    context_summary = st.session_state.get("context_summary_persisted", "")
    content_outline = st.session_state.get("content_outline", "")

    if not context_summary or not content_outline:
        st.error("Missing context summary or content outline. Please complete earlier steps.")
        return

    # Ask user if they want to generate the final assessment
    with st.form("assessment_form"):
        proceed = st.radio(
            "Do you want to generate the final assessment?",
            ("Yes, generate assessment", "No, finish here"),
            index=0
        )
        submitted = st.form_submit_button("Continue")
    
    if not submitted:
        st.stop()
    
        
    if proceed == "No, finish here":
        st.success("🎉 Instructional design process completed successfully without final assessment.")
        return

    

    num_questions = extract_num_questions(context_summary)
    if not num_questions:
        num_questions = estimate_questions_from_duration(context_summary)

    prompt = (
        f"Based on the following instructional design context and content outline, generate a final assessment for this e-learning course.\n\n"
        f"### Instructional Design Context:\n{context_summary}\n\n"
        f"### Content Outline:\n{content_outline}\n\n"
        f"Create {num_questions} multiple-choice questions.\n"
        f"Each MCQ must have appropriate number of answer options, and should clearly indicate the correct option.\n"
        f"Ensure questions align with the course objectives and learning content.\n"
        f"Do not add any explanation text or headings before or after the questions."
    )
    with st.spinner("Generating final assessment..."):
        assessment = get_openai_response(prompt, max_completion_tokens=4500)

    if assessment:
        st.markdown("### Final Assessment")
        st.write(assessment)

        if st.button("💾 Save Project"):
            try:
                from docx import Document
                from docx.shared import Pt
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                import os
                import re

                # Create output folder
                user_email = st.session_state.get("user_email", "unknown_user")
                folder = f"saved_projects/{user_email}"
                os.makedirs(folder, exist_ok=True)

                # Create doc
                doc = Document()
                doc.add_heading("Storyboard", level=1)

                # --- Storyboard Table ---
                lines = st.session_state.storyboard.strip().split("\n")
                headers = [h.strip() for h in lines[0].split("|") if h.strip()]
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Table Grid'

                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    cell = hdr_cells[i]
                    cell.text = header
                    for p in cell.paragraphs:
                        p.runs[0].bold = True

                for line in lines[2:]:
                    cols = [c.strip() for c in line.split("|") if c.strip()]
                    if len(cols) == len(headers):
                        row = table.add_row().cells
                        for i, text in enumerate(cols):
                            run = row[i].paragraphs[0].add_run(text)
                            run.font.size = Pt(10)

                doc.add_page_break()
                doc.add_heading("Final Assessment", level=1)
                doc.add_paragraph(assessment)

                # Save file
                from datetime import datetime
                ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"{folder}/{ts}_IDA_Project.docx"
                doc.save(filename)
                
                from db_manager import save_project_metadata
                save_project_metadata(user_email, st.session_state.get("project_title", "Untitled Project"), filename)

                st.success("✅ Project saved successfully!")

                with open(filename, "rb") as f:
                    st.download_button("📥 Download Combined Word File", f, file_name="IDA_Project.docx")

                st.session_state.step = None

            except Exception as e:
                st.error(f"❌ Failed to save project: {e}")

        
    else:
        st.error("Failed to generate final assessment. Please retry.")

def extract_num_questions(summary):
    match = re.search(r"(\d+)\s*(questions|mcqs|multiple choice)", summary.lower())
    if match:
        return int(match.group(1))
    return None

def estimate_questions_from_duration(summary):
    match = re.search(r"duration\s*[:\-]\s*(\d+)\s*(minutes|min|hours|hrs)", summary.lower())
    if match:
        duration = int(match.group(1))
        return max(1, duration // 10)  # Estimate 1 question per 10 minutes
        # unit = match.group(2)
        # if "hour" in unit:
        #     duration = duration * 60
        # if duration < 30:
        #     return 5
        # elif duration < 60:
        #     return 8
        # elif duration < 120:
        #     return 12
        # else:
        #     return 15
    return None

def generate_assessment(context_summary, content_outline):
    num_questions = extract_num_questions(context_summary)
    if not num_questions:
        num_questions = estimate_questions_from_duration(context_summary)

    assessment = create_final_assessment(context_summary, content_outline, num_questions)
    return assessment