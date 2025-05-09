import streamlit as st
import openai
import pandas as pd
import os
import io

# Set your OpenAI API key securely
# It's recommended to use environment variables or Streamlit secrets for API keys
openai.api_key = os.getenv("OPENAI_API_KEY")  # Ensure you set this environment variable
max_completion_tokens=10000

# Initialize session state variables
if 'step' not in st.session_state:
    st.session_state.step = 1
if 'context' not in st.session_state:
    st.session_state.context = {}
if 'raw_content' not in st.session_state:
    st.session_state.raw_content = None
if 'analysis' not in st.session_state:
    st.session_state.analysis = None
if 'content_outline' not in st.session_state:
    st.session_state.content_outline = None
if 'storyboard' not in st.session_state:
    st.session_state.storyboard = None
if 'final_assessment' not in st.session_state:
    st.session_state.final_assessment = None

# Function to call OpenAI API
def get_openai_response(prompt, max_completion_tokens=3500):
    try:
        response = openai.ChatCompletion.create(
            model="o1",  # Ensure you have access to GPT-4 or use "gpt-3.5-turbo"
            messages=[
                {"role": "system", "content": "You are a professional instructional design assistant."},
                {"role": "user", "content": prompt}
            ],
            max_completion_tokens=max_completion_tokens,
            n=1,  # Number of responses to generate
            stop=None  # Define stop sequences if needed
        )
        return response.choices[0].message['content'].strip()
#    except openai.error.OpenAIError as e:
#        st.error(f"OpenAI API returned an error: {e}")
#        return None
    except Exception as e:
        st.error(f"An unexpected error occurred: {e}")
        return None

# Step 1: Gather Context Information
def gather_context():
    st.header("Step 1: Gather Course Information")

    if "greeted" not in st.session_state:
        st.session_state.greeted = False
    if not st.session_state.greeted:
        st.write("👋 **Welcome!** How're you doing today? Seems like you are looking to create an e-learning course. Let's work together on it!")
        st.session_state.greeted = True

    if "context" not in st.session_state:
        st.session_state.context = {}
    if "conversation_history" not in st.session_state:
        st.session_state.conversation_history = [
            {"role": "system", "content": (
                "You are an instructional design assistant. Your goal is to collect all essential details such as the topic, audience profile (like age profile, education and experience), key learning outcomes, mode of learning (instructor lead or self paced), duration of the course, whether the user has raw content for the course, whether a final assessment and knowledge checks are needed and any other information the user has about the context, required to design an e-learning course. "
                "You can ask a maximum of 7 questions so prioritize the questions on its importance. Ask only one question at a time. If the response needs clarification, ask a follow-up which won't be counted in the 7 questions. "
                "When the user responds to a question, analyze it and see if there is a follow-up question needed to understand the context better. If yes, present that question. Don't ask too many questions but if the user's response is not clear, do not hesitate to ask." 
                "Once the answer is sufficient, move to the next key topic. Keep the conversation smooth and engaging."
            )}
        ]
    if "current_question" not in st.session_state:
        st.session_state.current_question = "What is the topic of your e-learning course?"
    if "question_count" not in st.session_state:
        st.session_state.question_count = 0
    if "context_complete" not in st.session_state:
        st.session_state.context_complete = False

    question_limit = 9  # 7 core + 2 follow-up questions

    if not st.session_state.context_complete:
        st.write(f"**{st.session_state.current_question}**")

        # 🔑 Dynamic key for text input to ensure clearing
        input_key = f"user_input_{len(st.session_state.conversation_history)}"
        user_input = st.text_area("Your Response:", key=input_key, height=150)

        if st.button("Submit"):
            if user_input.strip():
                st.session_state.context[st.session_state.current_question] = user_input
                st.session_state.conversation_history.append({"role": "user", "content": user_input})
                st.session_state.question_count += 1

                if st.session_state.question_count >= question_limit:
                    st.session_state.context_complete = True
                else:
                    prompt = (
                        "Here is the conversation so far:\n"
                        + "\n".join([f"{msg['role']}: {msg['content']}" for msg in st.session_state.conversation_history])
                        + "\n\nBased on the context gathered, ask the next most relevant question needed to design the course. "
                        "There is a limit of 6 questions, so prioritize the most important ones."
                    )
                    with st.spinner("Thinking what else do we need to know..."):
                        next_question = get_openai_response(prompt)
                        if next_question:
                            st.session_state.current_question = next_question
                            st.session_state.conversation_history.append({"role": "assistant", "content": next_question})
                        else:
                            st.session_state.context_complete = True

                # Force a rerun so the input box gets refreshed
                st.rerun()
            else:
                st.warning("Please enter your response before submitting.")

    if st.session_state.context_complete and "context_summary" not in st.session_state:
    # Generate summarized version with LLM only once
        summary_prompt = (
            "Summarize the following instructional design context into concise bullet points. "
            "Each bullet should have a short label (up to 6 words) representing the key aspect collected, "
            "and a short summary (not the full user response). Avoid repeating full questions or raw text. "
            "Return the result as a two-column table with headers 'Aspect' and 'Summary'.\n\n"
            f"Context:\n{st.session_state.context}"
        )
        with st.spinner("Summarizing context..."):
            summary_result = get_openai_response(summary_prompt)
            st.session_state.context_summary = summary_result
            st.session_state.context_summary_persisted = summary_result

    if st.session_state.context_complete and "context_summary" in st.session_state:
        st.subheader("✅ Summary of Collected Context")
        st.markdown(st.session_state.context_summary)

        cols = st.columns(2)
        with cols[0]:
            if st.button("Approve and Continue"):
                # 🔍 Check for raw content answer before proceeding to Step 2
                has_raw_content = None
                for question, answer in st.session_state.context.items():
                    if "raw content" in question.lower():
                        has_raw_content = "yes" in answer.lower()
                        break
                st.session_state.has_raw_content = has_raw_content

                st.session_state.step = 2
                # del st.session_state.context_summary
                st.rerun()
    
        with cols[1]:
            if st.button("Modify Information"):
                st.session_state.context_complete = False
                st.session_state.question_count = 0
                st.session_state.context = {}
                st.session_state.current_question = "What is the topic of your e-learning course?"
                st.session_state.conversation_history = [
                    {"role": "system", "content": st.session_state.conversation_history[0]["content"]}
                ]
                if "context_summary" in st.session_state:
                    del st.session_state.context_summary
                st.rerun()

                # Reset state
                st.session_state.context_complete = False
                st.session_state.question_count = 0
                st.session_state.context = {}
                st.session_state.current_question = "What is the topic of your e-learning course?"
                st.session_state.conversation_history = [
                    {"role": "system", "content": st.session_state.conversation_history[0]["content"]}
                ]
                st.rerun()

# Step 2: Analyze Raw Content
import re
from docx import Document
from io import BytesIO

def analyze_content():
    st.header("Step 2: Analyze Raw Content")

    uploaded_files = st.file_uploader(
        "Upload raw content files (PDF, DOCX, TXT)", 
        type=["pdf", "docx", "txt"], 
        accept_multiple_files=True
    )

    if uploaded_files:
        raw_text = ""
        try:
            for uploaded_file in uploaded_files:
                if uploaded_file.type == "application/pdf":
                    from PyPDF2 import PdfReader
                    reader = PdfReader(uploaded_file)
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            raw_text += page_text + "\n"
                elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    import docx
                    doc = docx.Document(uploaded_file)
                    for para in doc.paragraphs:
                        raw_text += para.text + "\n"
                elif uploaded_file.type == "text/plain":
                    raw_text += uploaded_file.getvalue().decode("utf-8") + "\n"
        except Exception as e:
            st.error(f"Error reading uploaded file: {e}")
            return

        if "analysis_done" not in st.session_state or st.session_state.raw_text != raw_text:
            st.session_state.raw_text = raw_text
            context_summary = st.session_state.get("context_summary", "No context summary available.")
            prompt = (
                f"Analyze the following instructional design context and raw content to identify content gaps.\n\n"
                f"Here is the instructional design context:\n{context_summary}\n\n"
                f"Here is the raw content:\n{raw_text}\n\n"
                f"Identify any content gaps in the raw content based on the provided context. Don't make it very elaborate and focus on the duration indicated by the user in {context_summary}. "
                f"List the missing topics or areas that need to be covered in the course."
            )
            with st.spinner("Analyzing content gaps..."):
                analysis = get_openai_response(prompt)
                st.session_state.analysis = analysis
                st.session_state.analysis_done = True
                st.session_state.uploaded_content = raw_text

        st.subheader("Content Gap Analysis")
        st.write(st.session_state.analysis)

        decision = st.radio(
            "How would you like to address the content gaps?",
            options=("Generate content to fill gaps", "Provide additional sources", "No action needed"),
            index=2
        )

        if decision == "Generate content to fill gaps":
            filled_prompt = (
                f"Based on the identified content gaps below, generate the necessary content to fill these gaps.\n\n"
                f"**Content Gaps:**\n{st.session_state.analysis}\n\n"
                f"Provide the additional content required to cover these areas effectively."
            )
            with st.spinner("Generating additional content..."):
                filled_content = get_openai_response(filled_prompt)
                if filled_content:
                    st.session_state.filled_content = filled_content
                    st.session_state.generated_additional_content = filled_content
                    st.success("Content gaps have been filled with generated material.")
                    if st.button("Continue to Step 3"):
                        st.session_state.step = 3
                        st.rerun()

        elif decision == "Provide additional sources":
            more_files = st.file_uploader(
                "Upload additional files to improve coverage",
                type=["pdf", "docx", "txt"],
                accept_multiple_files=True,
                key="more_sources"
            )
            if more_files:
                uploaded_files.extend(more_files)
                st.warning("Please click the 'Analyze Again' button below to include the new files.")
                if st.button("Analyze Again"):
                    st.rerun()

        elif decision == "No action needed":
            if st.button("Continue to Step 3"):
                st.session_state.step = 3
                st.rerun()
    else:
        st.info("Please upload at least one raw content file to begin analysis.")



# Step 3: Generate Content Outline
def generate_outline():
    st.header("Step 3: Generate Content Outline")

    context_summary = st.session_state.get("context_summary_persisted", "")
    uploaded_content = st.session_state.get("uploaded_content", "")
    generated_content = st.session_state.get("generated_additional_content", "")

    if not context_summary:
        st.error("Context summary not available. Please complete Step 1.")
        return

    if not uploaded_content and not generated_content:
        st.error("No content available. Please upload content or choose to generate content in Step 2.")
        return

    combined_content = uploaded_content.strip()
    if generated_content:
        combined_content += f"\n\n{generated_content.strip()}"

    if not st.session_state.get("content_outline"):
        prompt = (
            f"Based on the following instructional design context and source content, generate a structured content outline. Make sure that the duration indicated by the user in context_summary_persisted is adhered to when the duration in the outline is generated "
            f"for the e-learning course.\n\n"
            f"### Instructional Design Context:\n{context_summary}\n\n"
            f"### Source Content:\n{combined_content}\n\n"
            f"Present the content outline **strictly** as a table with two columns: Outline | Duration (in mins). Use pipe separators (|) and include a header row. Do not use bullets, dashes, or markdown separators.\n"
            f"Start immediately with the table header. Separate columns using a '|' (pipe symbol).\n"
            f"Do not add bullets or explanations before or after the table."
        )
        with st.spinner("Generating content outline..."):
            outline = get_openai_response(prompt)
            if outline:
                st.session_state.content_outline = outline

    if "content_outline" in st.session_state:
        st.subheader("Generated Content Outline")
        try:
            import pandas as pd
            import io
            df = pd.read_csv(io.StringIO(st.session_state.content_outline), sep="|")

            # Clean headers
            df.columns = [col.strip() for col in df.columns]

            # Drop any unnamed/empty columns
            df = df.loc[:, ~df.columns.str.contains("^Unnamed")]

            # Reset index to remove auto-numbering from display
            df = df.reset_index(drop=True)
            
            # Retain only the last two columns (in case extra sneaked in)
            if df.shape[1] > 2:
                df = df.iloc[:, -2:]

            # Rename headers to exactly what we want
            df.columns = ["Outline", "Duration (in mins)"]

            # Wrap long text inside cells
            st.markdown(
                """
                <style>
                .stDataFrame td {
                    white-space: pre-wrap !important;
                    word-break: break-word !important;
                }
                </style>
                """,
                unsafe_allow_html=True
            )

            # Display the table without index
            st.dataframe(df, use_container_width=True)

        except Exception as e:
            st.error("Could not parse outline as table. Showing raw text.")
            st.code(st.session_state.content_outline)

    with st.form("approve_outline_form"):
        proceed = st.form_submit_button("Approve Outline and Continue", type="primary")
        if proceed:
            st.session_state.step = 4
            st.rerun()



# Step 4: Generate Storyboard
def generate_storyboard():
    st.header("Step 4: Generate Storyboard")

    context_summary = st.session_state.get("context_summary_persisted", "")
    content_outline = st.session_state.get("content_outline", "")
    uploaded_content = st.session_state.get("uploaded_content", "")
    generated_content = st.session_state.get("generated_additional_content", "")

    if not content_outline:
        st.error("No content outline found. Please complete Step 3 before proceeding.")
        return

    combined_content = uploaded_content.strip()
    if generated_content:
        combined_content += f"\n\n{generated_content.strip()}"

    if not st.session_state.get("storyboard"):
        prompt = (
            f"Create a storyboard that follows instructional design theories, for the e-learning course based on the following instructional design context, content outline, and source content.\n\n"
            f"### Instructional Design Context:\n{context_summary}\n\n"
            f"### Content Outline:\n{content_outline}\n\n"
            f"### Source Content:\n{combined_content}\n\n"
            f"Provide the storyboard as a table with three columns: Onscreen Text | Voice Over Script | Visualization Guidelines.\n"
            f"Start immediately with the table header. Separate columns using a '|' (pipe symbol).\n"
            f"Make sure that the Onscreen text column contains the entire text we want to include in the slide, not just slide titles. In the Voice over script column, include the entire narrative voice over script, not just an introduction.\n"
            f"Let knowledge checks not be too many. Also when knowledge checks are used include all the details - the question, the answer options, the correct answer and also correct and wrong answer feedback. \n"
            f"Do not add any explanation before or after the table. Each row must be properly formatted without bullets or other formatting."
        )

        with st.spinner("Generating storyboard. This may take a while..."):
            storyboard = get_openai_response(prompt, max_completion_tokens=20000)
            if storyboard:
                st.session_state.storyboard = storyboard

    storyboard_text = st.session_state.get("storyboard", "")

    if storyboard_text:
        st.subheader("Generated Storyboard")
        import io
        import pandas as pd
        try:
            df = pd.read_csv(io.StringIO(st.session_state.storyboard), sep="|")
            df.columns = [col.strip() for col in df.columns]
            df = df.loc[:, ~df.columns.str.contains("^Unnamed")]
            df = df.reset_index(drop=True)

            st.markdown(
                """
                <style>
                .stDataFrame td {
                    white-space: pre-wrap !important;
                    word-break: break-word !important;
                }
                </style>
                """,
                unsafe_allow_html=True
            )

            st.dataframe(df, use_container_width=True)

        except Exception:
            st.error("Could not render storyboard as a table here. Showing raw text instead. Don't worry, export to word and it will be fine")
            st.code(st.session_state.storyboard)

            
        # Export to Word
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import io

        lines = storyboard_text.strip().split("\n")

        if len(lines) >= 2:
            buffer = io.BytesIO()
            doc = Document()
            doc.add_heading('Storyboard', level=1)

            headers = [h.strip() for h in lines[0].split("|") if h.strip()]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            table.autofit = True
            hdr_cells = table.rows[0].cells
            for idx, col in enumerate(headers):
                cell = hdr_cells[idx]
                cell.text = col
                for paragraph in cell.paragraphs:
                    run = paragraph.runs[0]
                    run.bold = True
                    run.font.size = Pt(11)

            for line in lines[2:]:
                columns = [c.strip() for c in line.split("|") if c.strip()]
                if len(columns) == len(headers):
                    row_cells = table.add_row().cells
                    for idx, val in enumerate(columns):
                        cell = row_cells[idx]
                        paragraph = cell.paragraphs[0]
                        run = paragraph.add_run(str(val))
                        run.font.size = Pt(10)

                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcW = OxmlElement('w:tcW')
                        tcW.set(qn('w:type'), 'auto')
                        tcW.set(qn('w:w'), '0')
                        tcPr.append(tcW)

            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="\U0001F4C4 Download Storyboard as Word file",
                data=buffer,
                file_name="Storyboard.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.error("Storyboard format not valid. Cannot export to Word.")
    else:
        st.error("Storyboard could not be generated. Please retry.")

    with st.form("approve_storyboard_form"):
        submitted = st.form_submit_button("Approve Storyboard and Continue", type="primary")
        if submitted:
            st.session_state.step = 5
            st.rerun()





# Step 5: Create Final Assessment
import re
from docx import Document
from io import BytesIO

def create_final_assessment():
    if st.session_state.get("assessment_downloaded"):
        st.success("🎉 Instructional design process completed successfully!")
        return
    
    st.header("Step 5: Create Final Assessment")

    context_summary = st.session_state.get("context_summary_persisted", "")
    content_outline = st.session_state.get("content_outline", "")

    if not context_summary or not content_outline:
        st.error("Missing context summary or content outline. Please complete earlier steps.")
        return

    # Ask user if they want to generate the final assessment
    with st.form("assessment_form"):
        proceed = st.radio(
            "Do you want to generate the final assessment?",
            ("Yes, generate assessment", "No, finish here"),
            index=0
        )
        submitted = st.form_submit_button("Continue")
    
    if not submitted:
        st.stop()
    
      
    if proceed == "No, finish here":
        st.success("🎉 Instructional design process completed successfully without final assessment.")
        return

    # Helper 2: Try extracting number of questions (from original code)
    def extract_num_questions(summary):
        match = re.search(r"(\\d+)\\s*(questions|mcqs|multiple choice)", summary.lower())
        if match:
            return int(match.group(1))
        else:
            return None

    # Helper 3: Estimate based on course duration (from original code)
    def estimate_questions_from_duration(summary):
        match = re.search(r"duration\\s*[:\\-]\\s*(\\d+)\\s*(minutes|min|hours|hrs)", summary.lower())
        if match:
            duration = int(match.group(1))
            unit = match.group(2)
            if "hour" in unit:
                duration = duration * 60
            if duration < 30:
                return 5
            elif duration < 60:
                return 8
            elif duration < 120:
                return 12
            else:
                return 15
        else:
            return 7  # fallback default

    num_questions = extract_num_questions(context_summary)
    if not num_questions:
        num_questions = estimate_questions_from_duration(context_summary)

    prompt = (
        f"Based on the following instructional design context and content outline, generate a final assessment for this e-learning course.\n\n"
        f"### Instructional Design Context:\n{context_summary}\n\n"
        f"### Content Outline:\n{content_outline}\n\n"
        f"Create {num_questions} multiple-choice questions.\n"
        f"Each MCQ must have appropriate number of answer options, and should clearly indicate the correct option.\n"
        f"Ensure questions align with the course objectives and learning content.\n"
        f"Do not add any explanation text or headings before or after the questions."
    )
    with st.spinner("Generating final assessment..."):
        assessment = get_openai_response(prompt, max_completion_tokens=4500)

    if assessment:
        st.markdown("### Final Assessment")
        st.write(assessment)

        # Export to Word document
        doc = Document()
        doc.add_heading('Final Assessment', 0)
        doc.add_paragraph(assessment)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        downloaded = st.download_button(
            label="📥 Download as Word Document",
            data=buffer,
            file_name="final_assessment.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
        if downloaded:
            st.success("🎉 Instructional design process completed successfully!")
            st.session_state.step = None
            st.rerun
        
    else:
        st.error("Failed to generate final assessment. Please retry.")
        st.error("Failed to generate final assessment. Please retry.")


# Main Application Flow
def main():
    st.title("Instructional Design Agent")
    st.write("An interactive agent to help you develop a storyboard for your e-learning course.")

    if st.session_state.step == 1:
        gather_context()
    elif st.session_state.step == 2:
        analyze_content()
    elif st.session_state.step == 3:
        generate_outline()
    elif st.session_state.step == 4:
        generate_storyboard()
    elif st.session_state.step == 5:
        create_final_assessment()
    else:
        st.write("All steps completed. Thank you for using the Instructional Design Agent!")

if __name__ == "__main__":
    main()
